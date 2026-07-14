#!/usr/bin/env python3
"""Populate problems_en.json and problems_vi.json from SQL seeds and docx files."""

import json
import re
from docx import Document

# Problem ID mappings from AI_TRANSLATION_INSTRUCTIONS.md Section 4
ANSWER_PROBLEMS = [
    ('perm_cat', 'Permutations of "CAT"', 'Hoán vị của từ "CAT"'),
    ('perm_book', 'Permutations of "BOOK"', 'Hoán vị của từ "BOOK"'),
    ('comb_committee_3of10', 'Choosing a Committee of 3', 'Chọn 3 người trong 10'),
    ('binom_coeff_x3_xy7', 'Coefficient of x³ in (x + y)⁷', 'Hệ số của x³ trong khai triển (x + y)⁷'),
    ('stars_bars_nonneg_5', 'Non-negative Solutions to x₁+x₂+x₃=5', 'Nghiệm nguyên không âm của phương trình x₁ + x₂ + x₃ = 5'),
    ('pigeonhole_birth_month', 'Birth Month Pigeonhole', 'Tháng sinh (Nguyên lý Dirichlet)'),
    ('arrange_4_books', 'Arranging 4 Distinct Books', 'Sắp xếp 4 cuốn sách phân biệt'),
    ('binom_coeff_x2y4_xy6', 'Coefficient of x²y⁴ in (x + y)⁶', 'Hệ số của x²y⁴ trong khai triển (x + y)⁶'),
    ('balls_boxes_no_empty', 'Balls into Boxes (No Box Empty)', 'Phân phối bóng không hộp trống'),
    ('coin_flips_6_heads', 'Coin Flips with Exactly 6 Heads', 'Tung xu đúng 6 lần mặt ngửa'),
    ('sum_binom_coeff_n5', 'Sum of Binomial Coefficients (n = 5)', 'Tổng hệ số nhị thức với n = 5'),
    ('positive_int_solutions_4var', 'Positive Integer Solutions (4 Variables)', 'Nghiệm nguyên dương của phương trình 4 biến'),
    ('perm_mississippi', 'Permutations of "MISSISSIPPI"', 'Hoán vị của từ "MISSISSIPPI"'),
    ('binom_coeff_x4_1x10', 'Coefficient of x⁴ in (1 + x)¹⁰', 'Hệ số của x⁴ trong khai triển (1 + x)¹⁰'),
    ('pigeonhole_mod7', 'Pigeonhole: Same Remainder mod 7', 'Số dư khi chia cho 7 (Nguyên lý Dirichlet)'),
    ('candies_min_one_each', 'Distributing Candies with a Minimum Each', 'Phát kẹo tối thiểu 1 viên mỗi đứa'),
    ('functions_4_objects_2_boxes', 'Functions from 4 Objects to 2 Boxes', '4 vật vào 2 hộp'),
    ('sum_coeff_2x_3y_5', 'Sum of Coefficients in (2x − 3y)⁵', 'Tổng hệ số trong khai triển (2x − 3y)⁵'),
    ('committee_2men_2women', 'Committee with Exactly 2 Men and 2 Women', 'Ban đại diện đúng 2 nam 2 nữ'),
    ('prove_symmetry_identity', 'Proving C(n,k) = C(n, n−k)', 'Chứng minh đẳng thức C(n,k) = C(n, n−k)'),
]

CODE_PROBLEMS = [
    ('sage_multiset_no_adj', 'Multiset Permutations, No Equal Letters Adjacent', 'Hoán vị đa tập hợp không chữ cái liền kề'),
    ('sage_circular_arrangements', 'Circular Arrangements (up to Rotation)', 'Sắp xếp vòng tròn (khác nhau theo phép quay)'),
    ('sage_circular_not_adjacent', 'Circular Arrangements with a Forbidden Adjacency', 'Sắp xếp vòng tròn với cặp liền kề bị cấm'),
    ('sage_multinomial_coeff', 'Multinomial Coefficient from an Expansion', 'Hệ số đa thức từ khai triển'),
    ('sage_bounded_stars_bars', 'Bounded Stars and Bars', 'Vách ngăn có giới hạn'),
    ('sage_surjections', 'Counting Surjections', 'Đếm toàn ánh'),
    ('sage_set_partitions_k_blocks', 'Set Partitions into Exactly k Blocks', 'Phân hoạch tập hợp thành đúng k khối'),
    ('sage_bell_number', 'Bell Numbers', 'Số Bell'),
    ('sage_gaussian_binomial', 'Gaussian Binomial Coefficients', 'Hệ số nhị thức Gauss'),
    ('sage_derangements_r_fixed', 'Permutations with Exactly r Fixed Points', 'Hoán vị với đúng r điểm cố định'),
    ('sage_poisson_binomial', 'Poisson-Binomial: Exact Probability of k Successes', 'Xác suất Poisson-Binomial đúng k lần thành công'),
    ('sage_forced_nonempty_boxes', 'Onto Functions with Forced Nonempty Boxes', 'Toàn ánh với hộp bị ràng buộc khác rỗng'),
    ('sage_mixed_gen_func_coeff', 'Coefficient from a Mixed Generating Function', 'Hệ số từ hàm sinh hỗn hợp'),
    ('sage_ordered_bell', 'Ordered Bell (Fubini) Numbers', 'Số Bell có thứ tự (Fubini)'),
    ('sage_exactly_r_distinct', 'Exactly r Distinct Values Among n Draws', 'Đúng r giá trị phân biệt trong n lần rút'),
    ('sage_restricted_compositions', 'Restricted Compositions via Generating Functions', 'Phân tích có ràng buộc bằng hàm sinh'),
    ('sage_perms_c_cycles', 'Permutations with a Given Number of Cycles', 'Hoán vị với số chu trình cho trước'),
    ('sage_binomial_tail', 'Binomial Tail Probability (Symbolic)', 'Xác suất đuôi nhị thức dạng biểu tượng'),
    ('sage_binary_necklaces', 'Binary Necklaces (Burnside\'s Lemma)', 'Vòng cổ nhị phân (Bổ đề Burnside)'),
    ('sage_necklaces_k_black', 'Binary Necklaces with a Fixed Number of Black Beads', 'Vòng cổ nhị phân với số hạt đen cố định'),
]

def parse_answer_docx():
    """Parse Bai_tap_To_hop_Chinh_hop_Nhi_thuc_Newton.docx"""
    doc = Document('Bai_tap_To_hop_Chinh_hop_Nhi_thuc_Newton.docx')
    problems = {}
    current_problem = None
    in_solution = False
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # Detect problem start (Bài X)
        match = re.match(r'^Bài (\d+)$', text)
        if match:
            problem_num = int(match.group(1)) - 1  # 0-indexed
            if problem_num < len(ANSWER_PROBLEMS):
                current_problem = ANSWER_PROBLEMS[problem_num][0]
                problems[current_problem] = {'title': '', 'statement': ''}
                in_solution = False
            continue
        
        elif current_problem:
            # Extract title (first line after Bài X, before Đáp số)
            if not problems[current_problem]['title'] and not text.startswith('Đáp số:'):
                problems[current_problem]['title'] = text
            # Extract statement (description + solution)
            else:
                # Include everything after the title as part of the statement
                if problems[current_problem]['statement']:
                    problems[current_problem]['statement'] += ' ' + text
                else:
                    problems[current_problem]['statement'] = text
    
    return problems

def parse_code_docx():
    """Parse Bai_tap_To_hop_Xac_suat_SageMath.docx"""
    doc = Document('Bai_tap_To_hop_Xac_suat_SageMath.docx')
    problems = {}
    current_problem = None
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # Detect problem start (Bài X:)
        match = re.match(r'^Bài (\d+):', text)
        if match:
            problem_num = int(match.group(1)) - 1  # 0-indexed
            if problem_num < len(CODE_PROBLEMS):
                current_problem = CODE_PROBLEMS[problem_num][0]
                problems[current_problem] = {'title': '', 'statement': ''}
                # Extract title from the same line (after "Bài X:")
                title_part = text.split(':', 1)[1].strip() if ':' in text else ''
                if title_part:
                    problems[current_problem]['title'] = title_part
            continue
        
        elif current_problem:
            # Extract statement (description + solution)
            # Include everything after the title as part of the statement
            if problems[current_problem]['statement']:
                problems[current_problem]['statement'] += ' ' + text
            else:
                problems[current_problem]['statement'] = text
    
    return problems

def convert_to_html(statement_text):
    """Convert plain text statement to HTML with LaTeX preservation."""
    # Replace LaTeX patterns: \(...\) and \[...\]
    # This is a simple conversion - in production you'd want more sophisticated parsing
    html = statement_text
    # Wrap in <p> tags
    html = f'<p>{html}</p>'
    return html

def extract_statement_from_sql(sql_file, problem_id):
    """Extract statement field from SQL file for a given problem ID."""
    with open(sql_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Find the INSERT statement for this problem ID
    # Pattern: id, title, xp, answer, statement, type
    # We need to extract the statement (5th field)
    pattern = rf"'{problem_id}', '[^']*', \d+, '[^']*', '(.*?)', '[^']*'"
    match = re.search(pattern, content, re.DOTALL)
    if match:
        return match.group(1)
    
    return None

def main():
    # Load existing JSON files
    with open('backend/i18n/problems_en.json', 'r', encoding='utf-8') as f:
        problems_en = json.load(f)
    
    with open('backend/i18n/problems_vi.json', 'r', encoding='utf-8') as f:
        problems_vi = json.load(f)
    
    # Populate problems_en.json from SQL files
    for pid, en_title, vi_title in ANSWER_PROBLEMS:
        if pid in problems_en:
            problems_en[pid]['title'] = en_title
            statement = extract_statement_from_sql('backend/insert_practice_permutations.sql', pid)
            if statement:
                problems_en[pid]['statement'] = statement
    
    for pid, en_title, vi_title in CODE_PROBLEMS:
        if pid in problems_en:
            problems_en[pid]['title'] = en_title
            statement = extract_statement_from_sql('backend/insert_sage_problems.sql', pid)
            if statement:
                problems_en[pid]['statement'] = statement
    
    # Parse docx files for Vietnamese
    answer_vi = parse_answer_docx()
    code_vi = parse_code_docx()
    
    # Populate problems_vi.json
    for pid, en_title, vi_title in ANSWER_PROBLEMS:
        if pid in answer_vi:
            problems_vi[pid]['title'] = answer_vi[pid]['title']
            problems_vi[pid]['statement'] = convert_to_html(answer_vi[pid]['statement'])
    
    for pid, en_title, vi_title in CODE_PROBLEMS:
        if pid in code_vi:
            problems_vi[pid]['title'] = code_vi[pid]['title']
            problems_vi[pid]['statement'] = convert_to_html(code_vi[pid]['statement'])
    
    # Save updated files
    with open('backend/i18n/problems_en.json', 'w', encoding='utf-8') as f:
        json.dump(problems_en, f, ensure_ascii=False, indent=2)
    
    with open('backend/i18n/problems_vi.json', 'w', encoding='utf-8') as f:
        json.dump(problems_vi, f, ensure_ascii=False, indent=2)
    
    print("Populated problems_en.json and problems_vi.json")
    print(f"Answer problems: {len(answer_vi)}")
    print(f"Code problems: {len(code_vi)}")

if __name__ == '__main__':
    main()
